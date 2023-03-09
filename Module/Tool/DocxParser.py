import docx

import Module.Tool.Filter as Filter


def openDocument(Path: str):
    return docx.Document(Path)


def compareGroups(MyGroup: str, OtherGroup: str):
    MyGroup    = list(MyGroup.lower())
    OtherGroup = list(OtherGroup.lower())

    return set(MyGroup).issubset(OtherGroup)


def getLastRow(Document: docx.Document, Index: int):
    LastIndex = len(Document.tables[0].rows[Index].cells) - 1
    return Document.tables[0].rows[Index].cells[LastIndex].text


def getLessonCount(Lesson: str):
    return Lesson.replace(" ", "").split(",")


def normalize(Message: str):
	return Message.replace("\n", " ")


def groupInRows(Document: docx.Document, MyGroup: str):
    for Table in Document.tables:
        for Index, Row in enumerate(Table.rows):
            Group = str(Document.tables[0].rows[Index].cells[0].text)

            if (compareGroups(MyGroup = MyGroup, OtherGroup = Group)):
                return True
    
    return False


def parceReplace(Document: docx.Document, MyGroup: str):
    ReplaceList = list()

    for Table in Document.tables:
        for Index, Row in enumerate(Table.rows):
            Group = str(Document.tables[0].rows[Index].cells[0].text)

            if (compareGroups(MyGroup = MyGroup, OtherGroup = Group)):
                Lesson      = Document.tables[0].rows[Index].cells[1].text
                LessonCount = getLessonCount(Lesson = Lesson)

                Item  = normalize(Document.tables[0].rows[Index].cells[4].text)
                Teacher = Filter.getTeacherInStr(Item) if (("(" in Item) and ((")" in Item))) else Filter.getTeacher(Item)
                Item = Filter.removeTeacher(Item)

                Point = normalize(getLastRow(Document, Index))

                for _Lesson in LessonCount:
                    try   : ReplaceList.append((int(_Lesson), Item, Point, Teacher))
                    except: pass

    return ReplaceList